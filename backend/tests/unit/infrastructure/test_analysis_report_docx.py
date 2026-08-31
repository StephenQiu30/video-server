from io import BytesIO
from zipfile import ZipFile, is_zipfile

from app.application.analysis import render_analysis_report_markdown
from app.infrastructure.analysis_report_docx import PythonDocxAnalysisReportRenderer
from docx import Document
from tests.unit.application.analysis.test_report import report_result


def test_docx_report_is_valid_and_uses_business_brief_geometry() -> None:
    markdown = render_analysis_report_markdown(report_result())
    content = PythonDocxAnalysisReportRenderer().render(markdown)

    assert content.startswith(b"PK")
    assert is_zipfile(BytesIO(content))
    document = Document(BytesIO(content))
    section = document.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11
    assert round(section.left_margin.inches, 2) == 1
    assert document.core_properties.title == "产品 [演示](https://invalid.example)"
    assert document.core_properties.subject == "Editorial video analysis"
    assert len(document.tables) == 1
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "先说结论：这支片子最值得看什么" in paragraphs
    assert "如果继续打磨，先做这几件事" in paragraphs

    with ZipFile(BytesIO(content)) as package:
        xml = package.read("word/document.xml").decode("utf-8")
        header_xml = package.read("word/header1.xml").decode("utf-8")
    assert 'w:w="9360"' in xml
    assert 'w:w="2040"' in xml
    assert 'w:w="120"' in xml
    assert 'w:fill="F2F4F7"' in xml
    assert "EDITORIAL BREAKDOWN" in header_xml
    assert "展示主要界面。" in xml
