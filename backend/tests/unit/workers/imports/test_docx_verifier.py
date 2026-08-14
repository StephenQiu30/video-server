from __future__ import annotations

import hashlib
import stat
from collections.abc import Callable
from io import BytesIO
from pathlib import Path
from uuid import UUID
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pytest
from app.application.import_execution import (
    ImportVerificationClaim,
    ImportVerificationRejected,
)
from app.domain.imports import ContentKind, ImportErrorCode, ImportSourceFormat
from app.workers.imports import (
    DocxScreenplayVerifier,
    DocxVerificationSettings,
    TextVerificationSettings,
)
from docx import Document

DOCUMENT_ID = UUID("eeeeeeee-eeee-4eee-8eee-eeeeeeeeeeee")


def _docx(*, body: bool = True) -> bytes:
    document = Document()
    if body:
        document.add_paragraph("INT. ROOM - DAY")
        document.add_paragraph("ALICE")
        document.add_paragraph("Hello there.")
        document.add_paragraph("外景 - 夜")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "小明"
        table.cell(0, 1).text = "你好。"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _claim(content: bytes) -> ImportVerificationClaim:
    return ImportVerificationClaim(
        resource_id=DOCUMENT_ID,
        content_kind=ContentKind.SCREENPLAY,
        source_format=ImportSourceFormat.DOCX,
        attempt=1,
        version=2,
        object_key=f"quarantine/screenplay/{DOCUMENT_ID}/1/source",
        declared_size_bytes=len(content),
        declared_sha256=hashlib.sha256(content).hexdigest(),
    )


def _settings(**overrides: int) -> DocxVerificationSettings:
    values = {
        "max_entries": 512,
        "max_entry_bytes": 16 * 1024**2,
        "max_uncompressed_bytes": 64 * 1024**2,
        "max_compression_ratio": 100,
    }
    values.update(overrides)
    return DocxVerificationSettings(
        text=TextVerificationSettings(
            max_size_bytes=2 * 1024**2,
            max_characters=10_000,
            max_line_characters=1000,
        ),
        **values,
    )


async def _verify(
    tmp_path: Path,
    content: bytes,
    settings: DocxVerificationSettings | None = None,
):
    workspace = tmp_path / "task-random"
    workspace.mkdir(exist_ok=True)
    source = workspace / "source"
    source.write_bytes(content)
    verifier = DocxScreenplayVerifier(tmp_path, settings or _settings())
    return await verifier(source, _claim(content))


def _rewrite(
    content: bytes,
    *,
    replacements: dict[str, bytes] | None = None,
    extras: tuple[tuple[str | ZipInfo, bytes], ...] = (),
) -> bytes:
    replacements = replacements or {}
    output = BytesIO()
    with (
        ZipFile(BytesIO(content)) as source,
        ZipFile(output, "w", compression=ZIP_DEFLATED) as target,
    ):
        for entry in source.infolist():
            payload = replacements.get(entry.filename, source.read(entry))
            target.writestr(entry, payload)
        for name, payload in extras:
            target.writestr(name, payload)
    return output.getvalue()


async def test_docx_extracts_only_body_and_table_text_into_shared_contract(
    tmp_path: Path,
) -> None:
    verified = await _verify(tmp_path, _docx())

    normalized = verified.normalized_path.read_text(encoding="utf-8")
    assert "INT. ROOM - DAY\nALICE\nHello there." in normalized
    assert "外景 - 夜\n小明\t你好。" in normalized
    assert verified.detected_language == "mixed"
    assert len(verified.scenes) == 2
    assert verified.quality_warnings == ()


def _external(content: bytes) -> bytes:
    relationships = (
        b'<?xml version="1.0" encoding="UTF-8"?>'
        b'<Relationships xmlns="http://schemas.openxmlformats.org/package/'
        b'2006/relationships"><Relationship Id="rId1" Type="x" '
        b'Target="https://example.invalid/x" TargetMode="External"/>'
        b"</Relationships>"
    )
    return _rewrite(content, replacements={"_rels/.rels": relationships})


def _symlink(content: bytes) -> bytes:
    entry = ZipInfo("word/media/link")
    entry.create_system = 3
    entry.external_attr = (stat.S_IFLNK | 0o777) << 16
    return _rewrite(content, extras=((entry, b"../../outside"),))


def _bomb(content: bytes) -> bytes:
    entry = ZipInfo("word/media/bomb.txt")
    entry.compress_type = ZIP_DEFLATED
    return _rewrite(content, extras=((entry, b"x" * 100_000),))


@pytest.mark.parametrize(
    "mutator",
    [
        pytest.param(
            lambda value: _rewrite(value, extras=(("../escape.xml", b"x"),)),
            id="path-traversal",
        ),
        pytest.param(
            lambda value: _rewrite(value, extras=(("word/vbaProject.bin", b"macro"),)),
            id="macro",
        ),
        pytest.param(_external, id="external-relationship"),
        pytest.param(_symlink, id="symbolic-link"),
        pytest.param(_bomb, id="compression-ratio"),
        pytest.param(
            lambda value: _rewrite(
                value, extras=(("WORD/document.xml", b"duplicate"),)
            ),
            id="case-insensitive-duplicate",
        ),
    ],
)
async def test_docx_active_archive_and_external_content_fail_closed(
    tmp_path: Path, mutator: Callable[[bytes], bytes]
) -> None:
    content = mutator(_docx())

    with pytest.raises(ImportVerificationRejected) as raised:
        await _verify(tmp_path, content)

    assert raised.value.code is ImportErrorCode.DOCUMENT_ARCHIVE_UNSAFE


@pytest.mark.parametrize(
    "settings",
    [
        pytest.param(_settings(max_entries=2), id="entry-count"),
        pytest.param(_settings(max_entry_bytes=100), id="entry-size"),
        pytest.param(_settings(max_uncompressed_bytes=100), id="total-size"),
    ],
)
async def test_docx_capacity_limits_fail_before_extraction(
    tmp_path: Path, settings: DocxVerificationSettings
) -> None:
    content = _docx()

    with pytest.raises(ImportVerificationRejected) as raised:
        await _verify(tmp_path, content, settings)

    assert raised.value.code is ImportErrorCode.DOCUMENT_ARCHIVE_UNSAFE


async def test_invalid_and_empty_docx_return_stable_errors(tmp_path: Path) -> None:
    invalid = b"not an OOXML package"
    with pytest.raises(ImportVerificationRejected) as archive_error:
        await _verify(tmp_path, invalid)
    assert archive_error.value.code is ImportErrorCode.DOCUMENT_ARCHIVE_UNSAFE

    empty = _docx(body=False)
    with pytest.raises(ImportVerificationRejected) as text_error:
        await _verify(tmp_path, empty)
    assert text_error.value.code is ImportErrorCode.DOCUMENT_TEXT_UNAVAILABLE


async def test_encrypted_office_container_returns_encrypted_error(
    tmp_path: Path,
) -> None:
    encrypted = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"encrypted-package"

    with pytest.raises(ImportVerificationRejected) as raised:
        await _verify(tmp_path, encrypted)

    assert raised.value.code is ImportErrorCode.DOCUMENT_ENCRYPTED
