from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import _ParagraphStyle
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from markdown_it import MarkdownIt
from markdown_it.token import Token

_BLUE = RGBColor(0x2E, 0x74, 0xB5)
_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
_BLACK = RGBColor(0x0A, 0x0A, 0x0A)
_MUTED = RGBColor(0x68, 0x68, 0x68)
_TABLE_FILL = "F2F4F7"
_CONTENT_WIDTH_DXA = 9_360
_TABLE_INDENT_DXA = 120


class PythonDocxAnalysisReportRenderer:
    """Render the canonical Markdown report as a Word document."""

    def render(self, markdown: str) -> bytes:
        if not markdown.strip() or len(markdown) > 1_000_000:
            raise ValueError("markdown report must be non-blank and bounded")
        document = Document()
        title = _add_markdown(document, markdown)
        _configure_document(document, title)
        output = BytesIO()
        document.save(output)
        return output.getvalue()


def _configure_document(document: DocumentObject, title: str) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, "Calibri", 11, _BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Title", 21, _BLACK, 0, 4),
        ("Subtitle", 13, _MUTED, 0, 12),
        ("Heading 1", 16, _BLUE, 12, 6),
        ("Heading 2", 13, _BLUE, 10, 4),
        ("Heading 3", 12, _DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        _set_style_font(style, "Calibri", size, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        _set_style_font(style, "Calibri", 11, _BLACK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    document.core_properties.title = title
    document.core_properties.subject = "AI visual analysis report"
    document.core_properties.author = "Video Server"
    _configure_header(section)
    _configure_footer(section)


def _configure_header(section: Section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)
    values = ("VIDEO ANALYSIS REPORT", "SHOT-BY-SHOT ANALYSIS")
    for index, (cell, value) in enumerate(
        zip(table.rows[0].cells, values, strict=True)
    ):
        cell.width = Inches(3.25)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        )
        run = cell_paragraph.add_run(value)
        _set_run_font(run, 8, _MUTED, bold=True)


def _configure_footer(section: Section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Generated analysis report · Page ")
    _set_run_font(run, 8, _MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _add_markdown(document: DocumentObject, markdown: str) -> str:
    tokens = (
        MarkdownIt("commonmark", {"html": False, "linkify": False})
        .enable("table")
        .parse(markdown)
    )
    title = "Analysis report"
    list_styles: list[str] = []
    first_paragraph = True
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type == "heading_open":
            inline = _next_inline(tokens, index)
            text = _inline_text(inline.children or [])
            level = int(token.tag[1])
            if level == 1:
                title = text
                paragraph = document.add_paragraph(style="Title")
            else:
                paragraph = document.add_heading(level=min(level - 1, 3))
            _render_inline(paragraph, inline.children or [])
            first_paragraph = False
        elif token.type == "paragraph_open":
            inline = _next_inline(tokens, index)
            style = list_styles[-1] if list_styles else None
            if first_paragraph and style is None:
                style = "Subtitle"
            paragraph = document.add_paragraph(style=style)
            _render_inline(paragraph, inline.children or [])
            first_paragraph = False
        elif token.type == "ordered_list_open":
            list_styles.append("List Number")
        elif token.type == "bullet_list_open":
            list_styles.append("List Bullet")
        elif token.type in {"ordered_list_close", "bullet_list_close"}:
            list_styles.pop()
        elif token.type == "table_open":
            rows, index = _table_rows(tokens, index)
            _add_markdown_table(document, rows)
        index += 1
    return title


def _next_inline(tokens: list[Token], index: int) -> Token:
    if index + 1 >= len(tokens) or tokens[index + 1].type != "inline":
        raise ValueError("markdown block is missing inline content")
    return tokens[index + 1]


def _inline_text(tokens: list[Token]) -> str:
    return "".join(
        token.content for token in tokens if token.type in {"text", "code_inline"}
    )


def _render_inline(paragraph: Paragraph, tokens: list[Token]) -> None:
    bold = False
    italic = False
    for token in tokens:
        if token.type == "strong_open":
            bold = True
        elif token.type == "strong_close":
            bold = False
        elif token.type == "em_open":
            italic = True
        elif token.type == "em_close":
            italic = False
        elif token.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
        elif token.type in {"text", "code_inline"}:
            run = paragraph.add_run(token.content)
            _set_run_font(run, 11, _BLACK, bold=bold)
            run.italic = italic
            if token.type == "code_inline":
                run.font.name = "Consolas"


def _table_rows(tokens: list[Token], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    index = start + 1
    while index < len(tokens) and tokens[index].type != "table_close":
        token = tokens[index]
        if token.type == "tr_open":
            row = []
        elif token.type == "inline" and row is not None:
            row.append(_inline_text(token.children or []))
        elif token.type == "tr_close" and row is not None:
            rows.append(row)
            row = None
        index += 1
    if not rows or any(len(row) != len(rows[0]) for row in rows):
        raise ValueError("markdown table has an invalid shape")
    return rows, index


def _add_markdown_table(document: DocumentObject, rows: list[list[str]]) -> None:
    headings = tuple(rows[0])
    widths = _table_widths(len(headings))
    table = document.add_table(rows=1, cols=len(headings))
    _style_table(table, widths, headings)
    for row in rows[1:]:
        _add_table_row(table, tuple(row), widths)


def _table_widths(columns: int) -> tuple[int, ...]:
    if columns == 7:
        return (900, 1_140, 720, 2_040, 1_500, 2_040, 1_020)
    base, remainder = divmod(_CONTENT_WIDTH_DXA, columns)
    return tuple(base + (1 if index < remainder else 0) for index in range(columns))


def _style_table(
    table: Table,
    widths: tuple[int, ...],
    headings: tuple[str, ...],
) -> None:
    if sum(widths) != _CONTENT_WIDTH_DXA:
        raise ValueError("table widths must fill the 6.5 inch content area")
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(_CONTENT_WIDTH_DXA))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(_TABLE_INDENT_DXA))
    properties.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)
    _set_table_grid(table, widths)
    _set_repeat_table_header(table.rows[0])
    for cell, cell_width, heading in zip(
        table.rows[0].cells, widths, headings, strict=True
    ):
        _set_cell_width(cell, cell_width)
        _set_cell_fill(cell, _TABLE_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(heading)
        _set_run_font(run, 9, _BLACK, bold=True)


def _add_table_row(
    table: Table, values: tuple[str, ...], widths: tuple[int, ...]
) -> None:
    row = table.add_row()  # type: ignore[no-untyped-call]
    for cell, cell_width, value in zip(row.cells, widths, values, strict=True):
        _set_cell_width(cell, cell_width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(value)
        _set_run_font(run, 9, _BLACK)


def _set_style_font(
    style: _ParagraphStyle, name: str, size: int, color: RGBColor
) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    properties = style._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)


def _set_run_font(
    run: Run,
    size: int,
    color: RGBColor,
    *,
    bold: bool = False,
) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")


def _set_cell_width(cell: _Cell, width_dxa: int) -> None:
    cell.width = Inches(width_dxa / 1_440)
    tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width_dxa))
    margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell._tc.get_or_add_tcPr().append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)


def _set_table_grid(table: Table, widths: tuple[int, ...]) -> None:
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)


def _set_cell_fill(cell: _Cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_repeat_table_header(row: _Row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _remove_table_borders(table: Table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    table._tbl.tblPr.append(borders)
